"""
Generate Kayako CIM Executive Summary DOCX
Data sourced exclusively from Kayako_Confidential_Information_Memorandum.xlsx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x39, 0x64)   # headings / table header
TEAL   = RGBColor(0x00, 0x70, 0x96)   # sub-headings
LIGHT  = RGBColor(0xDF, 0xE9, 0xF3)   # table-header fill
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
DARK   = RGBColor(0x26, 0x26, 0x26)
RED    = RGBColor(0xC0, 0x00, 0x00)
ALT    = RGBColor(0xF2, 0xF7, 0xFC)   # alternating row fill

# ── Raw CIM data ──────────────────────────────────────────────────────────────
YEARS        = [2019, 2020, 2021, 2022, 2023, 2024]
REVENUE      = [12_500_000, 15_800_000, 21_200_000, 28_400_000, 36_900_000, 45_500_000]
ARR          = [10_200_000, 13_500_000, 18_900_000, 25_600_000, 33_200_000, 41_800_000]
GROSS_MARGIN = [68, 70, 72, 74, 76, 77]
EBITDA       = [12, 15, 18, 20, 22, 24]
CUSTOMERS    = [1_200, 1_450, 1_800, 2_300, 2_900, 3_400]
ARPU         = [8_500, 9_300, 10_500, 11_100, 11_500, 12_300]


# ── Helpers ───────────────────────────────────────────────────────────────────
def fmt_usd(n):
    """Format as $X.XM or $X.XK."""
    if abs(n) >= 1_000_000:
        return f"${n/1_000_000:,.1f}M"
    return f"${n/1_000:,.0f}K"

def fmt_pct(n, delta=False):
    prefix = "+" if (delta and n > 0) else ""
    return f"{prefix}{n}%"

def yoy_growth(series):
    """Return YoY % change list (first entry is None)."""
    result = [None]
    for i in range(1, len(series)):
        g = (series[i] - series[i-1]) / series[i-1] * 100
        result.append(round(g, 1))
    return result

def cagr(start, end, years):
    return round(((end / start) ** (1 / years) - 1) * 100, 1)

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "BFBFBF"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def cell_para(cell, text, bold=False, color=DARK, size=9,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    run  = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color


# ── Document builder ──────────────────────────────────────────────────────────
def build_document():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.4)
        section.right_margin  = Cm(2.4)

    # ── Default body style ──
    style = doc.styles["Normal"]
    style.font.name  = "Calibri"
    style.font.size  = Pt(10)
    style.font.color.rgb = DARK

    # ══════════════════════════════════════════════════════════════════════════
    # COVER / TITLE BLOCK
    # ══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KAYAKO")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Confidential Information Memorandum")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = TEAL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Financial Executive Summary")
    run.font.size = Pt(13)
    run.font.color.rgb = DARK

    doc.add_paragraph()  # spacer

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}   |   STRICTLY CONFIDENTIAL")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.italic = True

    # Horizontal rule
    doc.add_paragraph()
    hr = doc.add_paragraph()
    hr_run = hr.add_run("─" * 95)
    hr_run.font.color.rgb = NAVY
    hr_run.font.size = Pt(8)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION HEADING helper
    # ══════════════════════════════════════════════════════════════════════════
    def section_heading(text, numbered=""):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(f"{numbered}  {text}".strip())
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = NAVY

        # underline bar
        bar = doc.add_paragraph()
        bar.paragraph_format.space_before = Pt(0)
        bar.paragraph_format.space_after  = Pt(4)
        bar_run = bar.add_run("─" * 95)
        bar_run.font.color.rgb = TEAL
        bar_run.font.size = Pt(7)

    def sub_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = TEAL

    def body(text, bullet=False):
        p = doc.add_paragraph(style="List Bullet" if bullet else "Normal")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        return p

    def risk_bullet(text, label="RISK"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_after  = Pt(3)
        lbl = p.add_run(f"[{label}]  ")
        lbl.bold = True
        lbl.font.color.rgb = RED
        lbl.font.size = Pt(10)
        txt = p.add_run(text)
        txt.font.size = Pt(10)
        txt.font.color.rgb = DARK

    # ══════════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    section_heading("Executive Overview", "1.")

    rev_cagr  = cagr(REVENUE[0], REVENUE[-1], 5)
    arr_cagr  = cagr(ARR[0],     ARR[-1],     5)
    cust_cagr = cagr(CUSTOMERS[0], CUSTOMERS[-1], 5)

    overview_text = (
        f"Kayako is a B2B SaaS customer-support platform that has delivered sustained, "
        f"high-quality growth over the 2019–2024 period. Total revenue expanded from "
        f"{fmt_usd(REVENUE[0])} in 2019 to {fmt_usd(REVENUE[-1])} in 2024, representing a "
        f"5-year CAGR of {rev_cagr}%. Annual Recurring Revenue (ARR) reached "
        f"{fmt_usd(ARR[-1])} at year-end 2024, reflecting a 5-year CAGR of {arr_cagr}%. "
        f"Gross margin expanded 9 percentage points (68% → 77%) and EBITDA margin doubled "
        f"from 12% to 24%, evidencing meaningful operating leverage. The customer base grew "
        f"from 1,200 to 3,400 accounts ({cust_cagr}% CAGR), while ARPU rose 45% to $12,300, "
        f"signalling successful upsell / expansion-revenue motions."
    )
    body(overview_text)

    # KPI snapshot table
    doc.add_paragraph()
    snap_labels = ["Metric", "2019", "2024", "Change"]
    snap_data   = [
        ("Total Revenue",    fmt_usd(REVENUE[0]),   fmt_usd(REVENUE[-1]),   f"+{cagr(REVENUE[0],REVENUE[-1],5)}% CAGR"),
        ("ARR",              fmt_usd(ARR[0]),        fmt_usd(ARR[-1]),       f"+{arr_cagr}% CAGR"),
        ("Gross Margin",     fmt_pct(GROSS_MARGIN[0]), fmt_pct(GROSS_MARGIN[-1]), "+9 pp"),
        ("EBITDA Margin",    fmt_pct(EBITDA[0]),     fmt_pct(EBITDA[-1]),    "+12 pp"),
        ("Customers",        "1,200",                "3,400",                "+183%"),
        ("ARPU",             fmt_usd(ARPU[0]),       fmt_usd(ARPU[-1]),      "+45%"),
    ]

    tbl = doc.add_table(rows=1 + len(snap_data), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # header
    for i, h in enumerate(snap_labels):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, NAVY)
        cell_para(c, h, bold=True, color=WHITE, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    for r_idx, row_data in enumerate(snap_data):
        row = tbl.rows[r_idx + 1]
        bg  = ALT if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            align = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            cell_para(cell, val, bold=(c_idx == 0), size=9, align=align)

    # ══════════════════════════════════════════════════════════════════════════
    # 2. REVENUE OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    section_heading("Revenue Overview", "2.")

    body(
        "Kayako's revenue model is predominantly subscription-based (SaaS), with ARR "
        "constituting the majority of recognised revenue in each reported year. The "
        "consistently high ARR-to-revenue ratio (≥81% in every year) confirms a "
        "highly recurring, predictable revenue base."
    )

    doc.add_paragraph()

    # Revenue & ARR table
    rev_yoy  = yoy_growth(REVENUE)
    arr_yoy  = yoy_growth(ARR)
    arr_rev  = [round(ARR[i] / REVENUE[i] * 100, 1) for i in range(len(YEARS))]

    rev_headers = ["Year", "Revenue", "YoY Growth", "ARR", "YoY Growth", "ARR / Revenue"]
    tbl2 = doc.add_table(rows=1 + len(YEARS), cols=6)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.style = "Table Grid"

    for i, h in enumerate(rev_headers):
        c = tbl2.rows[0].cells[i]
        set_cell_bg(c, NAVY)
        cell_para(c, h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    for r_idx in range(len(YEARS)):
        row = tbl2.rows[r_idx + 1]
        bg  = ALT if r_idx % 2 == 0 else WHITE
        vals = [
            str(YEARS[r_idx]),
            fmt_usd(REVENUE[r_idx]),
            ("—" if rev_yoy[r_idx] is None else f"+{rev_yoy[r_idx]}%"),
            fmt_usd(ARR[r_idx]),
            ("—" if arr_yoy[r_idx] is None else f"+{arr_yoy[r_idx]}%"),
            f"{arr_rev[r_idx]}%",
        ]
        for c_idx, val in enumerate(vals):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell_para(cell, val, bold=(c_idx == 0), size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    body(
        "Revenue growth has been consistently strong, accelerating from 26.4% YoY (2019–2020) "
        "to a peak of 34.2% (2020–2021) before normalising at ~23% (2023–2024). This tapering "
        "is expected at scale and does not indicate structural deterioration — the absolute "
        f"revenue increment in 2024 (~{fmt_usd(REVENUE[-1]-REVENUE[-2])}) exceeded that of "
        "any prior year."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 3. GROWTH TRENDS
    # ══════════════════════════════════════════════════════════════════════════
    section_heading("Growth Trends", "3.")

    sub_heading("3.1  Revenue & ARR Momentum")
    body(
        f"Five-year revenue CAGR: {rev_cagr}%  |  Five-year ARR CAGR: {arr_cagr}%. "
        "ARR has grown faster than total revenue in every year, indicating that the "
        "subscription mix is increasing as a share of the business — a positive signal "
        "for revenue quality and forward visibility."
    )

    sub_heading("3.2  Margin Expansion")
    body(
        "Gross margin improved from 68% (2019) to 77% (2024), a gain of 9 percentage points "
        "over five years. This is consistent with SaaS scale economics: infrastructure and "
        "support costs grow sub-linearly relative to revenue as the customer base matures. "
        "EBITDA margin doubled from 12% to 24%, reflecting disciplined operating cost control "
        "alongside top-line growth."
    )

    # Margin table
    doc.add_paragraph()
    marg_headers = ["Year", "Revenue", "Gross Margin", "Gross Profit", "EBITDA Margin", "EBITDA ($)"]
    tbl3 = doc.add_table(rows=1 + len(YEARS), cols=6)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = "Table Grid"

    for i, h in enumerate(marg_headers):
        c = tbl3.rows[0].cells[i]
        set_cell_bg(c, NAVY)
        cell_para(c, h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    for r_idx in range(len(YEARS)):
        row  = tbl3.rows[r_idx + 1]
        bg   = ALT if r_idx % 2 == 0 else WHITE
        rev  = REVENUE[r_idx]
        gp   = round(rev * GROSS_MARGIN[r_idx] / 100)
        ebit = round(rev * EBITDA[r_idx] / 100)
        vals = [
            str(YEARS[r_idx]),
            fmt_usd(rev),
            fmt_pct(GROSS_MARGIN[r_idx]),
            fmt_usd(gp),
            fmt_pct(EBITDA[r_idx]),
            fmt_usd(ebit),
        ]
        for c_idx, val in enumerate(vals):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell_para(cell, val, bold=(c_idx == 0), size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    sub_heading("3.3  Customer & ARPU Growth")
    doc.add_paragraph()

    cust_yoy = yoy_growth(CUSTOMERS)
    arpu_yoy = yoy_growth(ARPU)

    ca_headers = ["Year", "Customers", "YoY Growth", "ARPU", "YoY Growth", "Total ARR"]
    tbl4 = doc.add_table(rows=1 + len(YEARS), cols=6)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl4.style = "Table Grid"

    for i, h in enumerate(ca_headers):
        c = tbl4.rows[0].cells[i]
        set_cell_bg(c, NAVY)
        cell_para(c, h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    for r_idx in range(len(YEARS)):
        row = tbl4.rows[r_idx + 1]
        bg  = ALT if r_idx % 2 == 0 else WHITE
        vals = [
            str(YEARS[r_idx]),
            f"{CUSTOMERS[r_idx]:,}",
            ("—" if cust_yoy[r_idx] is None else f"+{cust_yoy[r_idx]}%"),
            fmt_usd(ARPU[r_idx]),
            ("—" if arpu_yoy[r_idx] is None else f"+{arpu_yoy[r_idx]}%"),
            fmt_usd(ARR[r_idx]),
        ]
        for c_idx, val in enumerate(vals):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell_para(cell, val, bold=(c_idx == 0), size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    body(
        "Both customer count and ARPU have grown simultaneously — a strong indicator that "
        "Kayako is both acquiring new logos and successfully expanding existing accounts. "
        "ARPU growth of 45% over five years suggests effective product upsell and "
        "tiered-pricing realisation."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 4. ARR & RECURRING REVENUE INSIGHTS
    # ══════════════════════════════════════════════════════════════════════════
    section_heading("ARR & Recurring Revenue Insights", "4.")

    arr_pct_2024 = round(ARR[-1] / REVENUE[-1] * 100, 1)
    body(
        f"ARR of {fmt_usd(ARR[-1])} represents {arr_pct_2024}% of 2024 total revenue, "
        "confirming that the vast majority of revenue is subscription-based and therefore "
        "highly predictable. Key ARR observations:"
    )
    body("ARR has exceeded 80% of total revenue in every year of the reporting period.", bullet=True)
    body(
        f"Implied ARR expansion per customer (net ARR ÷ net new customers) rose from "
        f"~${round((ARR[-1]-ARR[-2])/(CUSTOMERS[-1]-CUSTOMERS[-2])):,} in 2023–2024, "
        "suggesting healthy net revenue retention.",
        bullet=True
    )
    body(
        "The gap between ARR and total revenue (non-ARR revenue) has remained relatively "
        "stable in absolute terms (~$2.3M–$3.7M), likely representing professional services, "
        "implementation fees, or one-time charges.",
        bullet=True
    )
    body(
        "ARR CAGR of 32.5% outpaces revenue CAGR of 29.5%, confirming the business is "
        "improving its recurring-revenue mix year over year.",
        bullet=True
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 5. RISK ANALYSIS
    # ══════════════════════════════════════════════════════════════════════════
    section_heading("Risk Analysis & Key Concerns", "5.")

    body(
        "The following risks and areas of diligence concern are identified based on the "
        "financial data presented in the CIM. Acquirers should seek confirmatory data "
        "in additional diligence materials."
    )
    doc.add_paragraph()

    risk_bullet(
        "Revenue Growth Deceleration — YoY revenue growth has slowed from 34.2% "
        "(2020–2021) to 23.3% (2023–2024). While partially attributable to base-rate "
        "effects at scale, the trend warrants scrutiny of pipeline quality, competitive "
        "positioning, and market saturation."
    )
    risk_bullet(
        "Non-ARR Revenue Opacity — The CIM discloses total revenue and ARR but does not "
        "break out services, implementation, or one-time revenue. The $3.7M gap between "
        "total revenue and ARR in 2024 may mask declining transactional revenue or "
        "one-time items that flatter headline figures."
    )
    risk_bullet(
        "Customer Concentration — No customer-level revenue breakdown is included in "
        "the CIM's Financial Summary sheet. With 3,400 customers and $45.5M in revenue, "
        "average spend is $12,300, but concentration in a small number of enterprise "
        "accounts cannot be ruled out without further analysis."
    )
    risk_bullet(
        "Churn & Net Revenue Retention — The CIM does not disclose gross or net revenue "
        "retention (NRR/GRR) rates. These are critical SaaS valuation metrics. Inferred "
        "ARR growth per existing customer is positive, but logo churn rate is unknown.",
        label="RISK"
    )
    risk_bullet(
        "EBITDA Quality — EBITDA margins of 24% are strong for a SaaS company of this "
        "scale, but the CIM does not clarify add-backs, stock-based compensation treatment, "
        "or capitalised software development costs. Adjusted vs. GAAP EBITDA reconciliation "
        "is required.",
        label="RISK"
    )
    risk_bullet(
        "Limited Segment Disclosure — No product-line, geographic, or vertical revenue "
        "breakdown is provided. Diversification across segments cannot be assessed from "
        "the current data.",
        label="INFO GAP"
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 6. KEY DILIGENCE QUESTIONS
    # ══════════════════════════════════════════════════════════════════════════
    section_heading("Key Diligence Questions for Management", "6.")

    questions = [
        "What is the gross and net revenue retention rate for 2022–2024?",
        "What is the revenue breakdown by product tier / module and by geography?",
        "What is the Top 10 / Top 20 customer revenue concentration as a % of ARR?",
        "What specific items constitute the non-ARR revenue of ~$3.7M in 2024?",
        "What is the EBITDA bridge from GAAP operating income to reported EBITDA "
        "(SBC, D&A, one-time items)?",
        "What is the average contract length and mix of monthly vs. annual subscriptions?",
        "What are the primary drivers of gross margin improvement and how sustainable are they?",
        "What is the expected ARR growth rate for FY2025 and what are the pipeline assumptions?",
    ]
    for i, q in enumerate(questions, 1):
        body(f"{i}.  {q}", bullet=False)

    # ══════════════════════════════════════════════════════════════════════════
    # 7. FINANCIAL SUMMARY APPENDIX
    # ══════════════════════════════════════════════════════════════════════════
    section_heading("Appendix: Full Financial Summary (2019–2024)", "7.")

    full_headers = [
        "Year", "Revenue ($)", "ARR ($)", "Gross Margin", "Gross Profit ($)",
        "EBITDA Margin", "EBITDA ($)", "Customers", "ARPU ($)"
    ]
    tbl5 = doc.add_table(rows=1 + len(YEARS), cols=len(full_headers))
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl5.style = "Table Grid"

    for i, h in enumerate(full_headers):
        c = tbl5.rows[0].cells[i]
        set_cell_bg(c, NAVY)
        cell_para(c, h, bold=True, color=WHITE, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    for r_idx in range(len(YEARS)):
        row  = tbl5.rows[r_idx + 1]
        bg   = ALT if r_idx % 2 == 0 else WHITE
        rev  = REVENUE[r_idx]
        gp   = round(rev * GROSS_MARGIN[r_idx] / 100)
        ebit = round(rev * EBITDA[r_idx] / 100)
        vals = [
            str(YEARS[r_idx]),
            fmt_usd(rev),
            fmt_usd(ARR[r_idx]),
            fmt_pct(GROSS_MARGIN[r_idx]),
            fmt_usd(gp),
            fmt_pct(EBITDA[r_idx]),
            fmt_usd(ebit),
            f"{CUSTOMERS[r_idx]:,}",
            fmt_usd(ARPU[r_idx]),
        ]
        for c_idx, val in enumerate(vals):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell_para(cell, val, bold=(c_idx == 0), size=8,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Footer note ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This document is based solely on information contained in the "
        "Kayako Confidential Information Memorandum (Financial_Summary sheet). "
        "All figures in USD. For internal M&A diligence use only."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    return doc


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    out_path = r"c:\github\Kayako-Due-Deligence\Kayako_CIM_Executive_Summary.docx"
    doc = build_document()
    doc.save(out_path)
    print(f"Saved: {out_path}")
